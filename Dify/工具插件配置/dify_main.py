'''
import json, re, uuid, os
from fastapi import FastAPI, Request
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from docx import Document

app = FastAPI()

# 允许跨域
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# 准备静态文件夹
if not os.path.exists("static"):
    os.makedirs("static")
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.post("/generate_word")
async def generate_word(request: Request):
    try:
        # 1. 拿原始数据（不报错的核心）
        body_bytes = await request.body()
        body_str = body_bytes.decode("utf-8")
        
        # 打印一下看看这坨“大便”模型发了什么
        print(f"--- 原始报文 --- \n{body_str} \n----------------")

        # 2. 暴力提取内容：寻找 "content": " 之后的所有文字
        # 这种正则哪怕 JSON 没闭合也能抓到内容
        content = ""
        pattern = r'"content":\s*"(.*)"'
        match = re.search(pattern, body_str, re.DOTALL)
        
        if match:
            content = match.group(1)
        else:
            # 保底方案：如果正则没抓到，尝试解析 JSON
            try:
                data = json.loads(body_str)
                content = data.get("content", body_str)
            except:
                content = body_str

        # 3. 关键：把 Unicode 转义字符（\u4f01）和 换行符（\\n）变回中文
        content = content.encode('utf-8').decode('unicode_escape', errors='ignore')
        # 去掉结尾可能多出的 JSON 残留引号
        content = content.rstrip('"}')

        # 4. 生成 Word
        doc = Document()
        doc.add_heading("TongWeb 8 产品分析报告", 0)
        for line in content.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
        
        file_name = f"{uuid.uuid4()}.docx"
        file_path = f"static/{file_name}"
        doc.save(file_path)
        
        return {"download_url": f"http://192.168.58.187:8000/static/{file_name}"}

    except Exception as e:
        print(f"报错: {e}")
        return {"error": "处理失败，但链接已生成"}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
'''

import uuid, os, re
from fastapi import FastAPI, Request, Form
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from docx import Document

app = FastAPI()

# 允许跨域
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# 准备静态文件夹（8000端口专门存word）
if not os.path.exists("static"):
    os.makedirs("static")
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.post("/generate_word")
async def generate_word(content: str = Form(...)): # 关键：用Form接收参数content
    try:
        print(f"--- 8000端口 收到 Word 请求 (长度: {len(content)}) ---")

        # 1. 清洗转义符
        # 有时候AI传过来的内容里会有字面量 \n，需要转回真正的换行
        clean_content = content.replace('\\n', '\n').replace('\\"', '"')

        # 2. 创建 Word 文档
        doc = Document()
        
        # 3. 智能排版逻辑
        lines = clean_content.splitlines()
        doc.add_heading("技术分析报告", 0) # 默认大标题
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 去掉 Markdown 样式的 # 和 *
            text = line.replace("#", "").replace("*", "").strip()
            
            # 正则识别：如果是以数字（如1.1）或者中文序号（如一、）开头，就加粗
            is_heading = re.match(r'^(\d+\.?\d*|[一二三四五六七八九十]+[、.])', text)
            
            if is_heading:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True # 变成加粗的小标题
            else:
                doc.add_paragraph(text)
        
        # 4. 保存文件
        file_name = f"{uuid.uuid4()}.docx"
        file_path = f"static/{file_name}"
        doc.save(file_path)
        
        # 5. 返回下载专用的 URL（走下面的 /download 接口）
        return {"download_url": f"http://192.168.137.52:8000/download/{file_name}"}
  
    except Exception as e:
        print(f"Word 插件内部报错: {e}")
        return {"error": str(e)}

# --- 新增：强制下载接口 ---
@app.get("/download/{file_name}")
async def download_file(file_name: str):
    file_path = f"static/{file_name}"
    if os.path.exists(file_path):
        # 这样返回会强制浏览器弹出“另存为”，而不是在网页里乱码打开
        return FileResponse(
            path=file_path, 
            filename=f"分析报告_{file_name[:8]}.docx", 
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    return {"error": "文件不存在或已过期"}

if __name__ == "__main__":
    import uvicorn
    # 坚守 8000 端口
    uvicorn.run(app, host="0.0.0.0", port=8000)